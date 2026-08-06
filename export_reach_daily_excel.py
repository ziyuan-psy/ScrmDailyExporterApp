#!/usr/bin/env python3
"""Write daily SCRM reach summaries to the app workbook.

The group reach calculation matches the confirmed business rule:
- delivered rows only
- match task "customer group chatid" to the daily customer-group base export
- sum group customer totals by task delivery row
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Sequence, Tuple
from xml.etree import ElementTree as ET

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo


NS_MAIN = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
NS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_PACKAGE_REL = "http://schemas.openxmlformats.org/package/2006/relationships"

FOLDER_PREFIX = "\u793e\u7fa4\u4efb\u52a1"
PRIVATE_ROOT_NAME = "\u6bcf\u65e5\u4f01\u5fae\u79c1\u57df\u4efb\u52a1\u5bfc\u51fa"
DEFAULT_OUTPUT_NAME = "\u793e\u7fa4\u4efb\u52a1\u89e6\u8fbe\u4eba\u6570\u65e5\u62a5.xlsx"
DEFAULT_REACH_DOCX_NAME = "\u4f01\u5fae\u793e\u7fa4\u4efb\u52a1\u89e6\u8fbe\u5ba2\u6237\u7edf\u8ba1.docx"

BASE_CHAT_ID_HEADER = "\u5ba2\u6237\u7fa4ID"
GROUP_CUSTOMER_TOTAL_HEADER = "\u7fa4\u5ba2\u6237\u603b\u6570"
TASK_CHAT_ID_HEADER = "\u5ba2\u6237\u7fa4chatid\uff08\u672c\u5e94\u7528\uff09"
DELIVERY_STATUS_HEADER = "\u9001\u8fbe\u72b6\u6001"
DELIVERED_STATUS = "\u5df2\u9001\u8fbe"
GROUP_SEND_NAME_HEADER = "\u5ba2\u6237\u7fa4\u540d"
SUPER_SEND_NAME_HEADER = "\u5ba2\u6237\u7fa4\u540d\u79f0"

SHEET_REACH_SUMMARY = "\u89e6\u8fbe\u4eba\u6570\u6c47\u603b"
SHEET_STORE_GROUP = "\u95e8\u5e97\u5206\u7ec4\u89e6\u8fbe\u4eba\u6570"
SHEET_README = "\u8bf4\u660e"

COL_DATE = "\u65e5\u671f"
COL_WEEKDAY = "\u661f\u671f"
COL_WELFARE_FRIENDS = "\u798f\u5229\u5b98\u597d\u53cb"
COL_COMMUNITY = "\u793e\u7fa4"
COL_MOMENTS = "\u670b\u53cb\u5708"
COL_SUPER_REACH = "\u8d85\u7ea7\u7fa4\u53d1\u89e6\u8fbe\u4eba\u6570"
COL_GROUP_REACH = "\u7fa4\u53d1\u5ba2\u6237\u7fa4\u89e6\u8fbe\u4eba\u6570"
COL_TOTAL_REACH = "\u5408\u8ba1\u89e6\u8fbe\u4eba\u6570"
COL_SUPER_ROWS = "\u8d85\u7ea7\u7fa4\u53d1\u5df2\u9001\u8fbe\u884c\u6570"
COL_GROUP_ROWS = "\u7fa4\u53d1\u5ba2\u6237\u7fa4\u5df2\u9001\u8fbe\u884c\u6570"
COL_MISSING = "\u672a\u5339\u914dchatid\u6570"
COL_FILE_COUNT = "\u4efb\u52a1\u6587\u4ef6\u6570"
COL_UPDATED_AT = "\u66f4\u65b0\u65f6\u95f4"
COL_FOLDER = "\u6765\u6e90\u6587\u4ef6\u5939"

REACH_SUMMARY_HEADERS = [
    COL_DATE,
    COL_WEEKDAY,
    COL_WELFARE_FRIENDS,
    "",
    "",
    COL_COMMUNITY,
    "",
    "",
    COL_MOMENTS,
]

COL_TYPE = "\u7c7b\u578b"
COL_STORE_GROUP = "\u95e8\u5e97\u5206\u7ec4"
COL_COUPON_TYPE = "\u5238\u7c7b\u578b"
COL_REACH = "\u89e6\u8fbe\u4eba\u6570"

STORE_GROUP_HEADERS = [
    COL_DATE,
    COL_TYPE,
    COL_STORE_GROUP,
    COL_COUPON_TYPE,
    COL_REACH,
]

STORE_GROUP_TYPE = "\u793e\u7fa4-\u5408\u8ba1"
ELASTIC_COUPON_TYPE = "\u5f39\u6027\u5238"
FOOD_COUPON_TYPE = "\u98df\u54c1\u5238"
STORE_GROUP_MARKERS = [
    ("A", "\u793e\u7fa4-A\u6863"),
    ("B", "\u793e\u7fa4-B\u6863"),
    ("C", "\u793e\u7fa4-C\u6863"),
    ("S", "\u793e\u7fa4-S\u6863"),
]
STORE_GROUP_MARKER_RE = re.compile(r"([ABCS])FD(?=_)")
DATE_HEADING_RE = re.compile(r"^(\d{4})\.(\d{1,2})\.(\d{1,2})$")
TOTAL_REACH_RE = re.compile(r"\u5171\u89e6\u8fbe([\d,]+)\u4eba")
TASK_REACH_RE = re.compile(r"\u4efb\u52a1\d+\u89e6\u8fbe([\d,]+)\u4eba")
GROUP_SEND_LABEL = "\u7fa4\u53d1\u5ba2\u6237"
MOMENT_SEND_LABEL = "\u7fa4\u53d1\u670b\u53cb\u5708"
WEEKDAYS = ["\u661f\u671f\u4e00", "\u661f\u671f\u4e8c", "\u661f\u671f\u4e09", "\u661f\u671f\u56db", "\u661f\u671f\u4e94", "\u661f\u671f\u516d", "\u661f\u671f\u65e5"]


class ReachSummaryError(RuntimeError):
    pass


@dataclass(frozen=True)
class DailyFolder:
    target_date: date
    path: Path


@dataclass
class GroupReachRow:
    target_date: date
    super_reach: int = 0
    group_reach: int = 0
    super_rows: int = 0
    group_rows: int = 0
    missing_chatid_count: int = 0
    task_file_count: int = 0
    updated_at: datetime = datetime.min
    folder_name: str = ""

    @property
    def total_reach(self) -> int:
        return self.super_reach + self.group_reach


@dataclass
class ReachSummaryRow:
    target_date: date
    welfare_friends: int = 0
    community: int = 0
    moments: int = 0

    def as_values(self) -> List[Any]:
        return [
            self.target_date,
            WEEKDAYS[self.target_date.weekday()],
            self.welfare_friends,
            None,
            None,
            self.community,
            None,
            None,
            self.moments,
        ]


@dataclass
class StoreGroupRow:
    target_date: date
    report_type: str
    store_group: str
    coupon_type: str
    reach: int

    def as_values(self) -> List[Any]:
        return [
            self.target_date,
            self.report_type,
            self.store_group,
            self.coupon_type,
            self.reach,
        ]


@dataclass
class DailyCalculation:
    group_reach: GroupReachRow
    store_group_rows: List[StoreGroupRow]


def parse_args(argv: Optional[Sequence[str]] = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Write SCRM daily reach summary workbook.")
    parser.add_argument("--source-root", default=None, help="Parent folder containing daily folders.")
    parser.add_argument("--output-xlsx", default=None, help="Output summary workbook path.")
    parser.add_argument("--reach-docx", default=None, help="Reach customer summary docx path.")
    parser.add_argument("--recent-days", type=int, default=5, help="Latest date folders to recompute. Default: 5.")
    parser.add_argument("--all", action="store_true", help="Recompute all date folders.")
    parser.add_argument("--date", action="append", default=[], help="Specific date to recompute, YYYY-MM-DD. Repeatable.")
    parser.add_argument("--folder-prefix", default=FOLDER_PREFIX, help="Daily folder prefix.")
    parser.add_argument("--today", default=None, help="Override current date for MMDD parsing, YYYY-MM-DD.")
    parser.add_argument("--summary-only", action="store_true", help=f"Only update {SHEET_REACH_SUMMARY}.")
    parser.add_argument("--store-group-only", action="store_true", help=f"Only update {SHEET_STORE_GROUP}.")
    parser.add_argument("--dry-run", action="store_true", help="Print results without writing the workbook.")
    return parser.parse_args(argv)


def default_source_root() -> Path:
    return Path.home() / "Documents" / PRIVATE_ROOT_NAME


def parse_today(value: Optional[str]) -> date:
    if value:
        return datetime.strptime(value, "%Y-%m-%d").date()
    return date.today()


def parse_folder_date(name: str, prefix: str, today: date) -> Optional[date]:
    if not name.startswith(prefix):
        return None
    match = re.search(r"(\d{2})(\d{2})$", name)
    if not match:
        return None
    month = int(match.group(1))
    day = int(match.group(2))
    try:
        parsed = date(today.year, month, day)
    except ValueError:
        return None
    if parsed > today:
        try:
            parsed = date(today.year - 1, month, day)
        except ValueError:
            return None
    return parsed


def find_daily_folders(source_root: Path, prefix: str, today: date) -> List[DailyFolder]:
    root_date = parse_folder_date(source_root.name, prefix, today)
    if root_date:
        return [DailyFolder(root_date, source_root)]

    if not source_root.exists():
        raise ReachSummaryError(f"Source root does not exist: {source_root}")

    folders: List[DailyFolder] = []
    for child in source_root.iterdir():
        if not child.is_dir():
            continue
        parsed = parse_folder_date(child.name, prefix, today)
        if parsed:
            folders.append(DailyFolder(parsed, child))
    return sorted(folders, key=lambda item: item.target_date)


def ns(tag: str) -> str:
    return f"{{{NS_MAIN}}}{tag}"


def package_ns(tag: str) -> str:
    return f"{{{NS_PACKAGE_REL}}}{tag}"


def first_sheet_path(zip_file: zipfile.ZipFile) -> str:
    try:
        workbook = ET.fromstring(zip_file.read("xl/workbook.xml"))
        sheets = workbook.find(ns("sheets"))
        first_sheet = sheets.find(ns("sheet")) if sheets is not None else None
        rel_id = first_sheet.attrib.get(f"{{{NS_REL}}}id") if first_sheet is not None else None
        rels = ET.fromstring(zip_file.read("xl/_rels/workbook.xml.rels"))
        for rel in rels.findall(package_ns("Relationship")):
            if rel.attrib.get("Id") == rel_id:
                target = rel.attrib.get("Target", "worksheets/sheet1.xml")
                return "xl/" + target.lstrip("/")
    except Exception:
        pass
    return "xl/worksheets/sheet1.xml"


def load_shared_strings(zip_file: zipfile.ZipFile) -> List[str]:
    try:
        with zip_file.open("xl/sharedStrings.xml") as handle:
            values: List[str] = []
            for event, element in ET.iterparse(handle, events=("end",)):
                if element.tag == ns("si"):
                    texts = [node.text or "" for node in element.iter(ns("t"))]
                    values.append("".join(texts))
                    element.clear()
            return values
    except KeyError:
        return []


def column_index_from_ref(ref: str) -> int:
    letters = []
    for char in ref:
        if char.isalpha():
            letters.append(char.upper())
        else:
            break
    result = 0
    for char in letters:
        result = result * 26 + ord(char) - ord("A") + 1
    return result


def cell_text(cell: ET.Element, shared_strings: List[str]) -> str:
    cell_type = cell.attrib.get("t", "")
    if cell_type == "inlineStr":
        return "".join(node.text or "" for node in cell.iter(ns("t")))

    value = cell.find(ns("v"))
    if value is None or value.text is None:
        return ""
    raw = value.text
    if cell_type == "s":
        try:
            return shared_strings[int(raw)]
        except (ValueError, IndexError):
            return ""
    return raw


def iter_sheet_rows(path: Path) -> Iterator[Dict[int, str]]:
    with zipfile.ZipFile(path) as zip_file:
        shared_strings = load_shared_strings(zip_file)
        sheet_path = first_sheet_path(zip_file)
        with zip_file.open(sheet_path) as handle:
            for event, element in ET.iterparse(handle, events=("end",)):
                if element.tag != ns("row"):
                    continue
                values: Dict[int, str] = {}
                next_column = 1
                for cell in element.findall(ns("c")):
                    ref = cell.attrib.get("r", "")
                    column = column_index_from_ref(ref) if ref else next_column
                    next_column = column + 1
                    values[column] = cell_text(cell, shared_strings).strip()
                yield values
                element.clear()


def row_value(row: Dict[int, str], column: int) -> str:
    return row.get(column, "").strip()


def classify_header(header: Dict[int, str]) -> Optional[str]:
    if row_value(header, 3) == BASE_CHAT_ID_HEADER and row_value(header, 10) == GROUP_CUSTOMER_TOTAL_HEADER:
        return "base"
    if row_value(header, 2) == TASK_CHAT_ID_HEADER and row_value(header, 8) == DELIVERY_STATUS_HEADER:
        if row_value(header, 1) == GROUP_SEND_NAME_HEADER:
            return "group_send"
        if row_value(header, 1) == SUPER_SEND_NAME_HEADER:
            return "super_send"
    return None


def to_int(value: str) -> int:
    if not value:
        return 0
    try:
        return int(float(value.replace(",", "")))
    except ValueError:
        return 0


def scan_workbook(path: Path) -> Tuple[Optional[str], Dict[str, int], List[str]]:
    rows = iter_sheet_rows(path)
    try:
        header = next(rows)
    except StopIteration:
        return None, {}, []

    kind = classify_header(header)
    if kind == "base":
        base_counts: Dict[str, int] = {}
        for row in rows:
            chat_id = row_value(row, 3)
            if chat_id:
                base_counts[chat_id] = to_int(row_value(row, 10))
        return kind, base_counts, []

    if kind in {"super_send", "group_send"}:
        delivered_chatids: List[str] = []
        for row in rows:
            if row_value(row, 8) == DELIVERED_STATUS:
                chat_id = row_value(row, 2)
                if chat_id:
                    delivered_chatids.append(chat_id)
        return kind, {}, delivered_chatids

    return kind, {}, []


def store_group_marker_from_name(name: str) -> Optional[str]:
    match = STORE_GROUP_MARKER_RE.search(name)
    return match.group(1) if match else None


def calculate_daily_folder(folder: DailyFolder) -> DailyCalculation:
    base_counts: Dict[str, int] = {}
    delivered_by_type: Dict[str, List[str]] = {"super_send": [], "group_send": []}
    store_group_delivered: Dict[str, List[str]] = {marker: [] for marker, _label in STORE_GROUP_MARKERS}
    task_file_count = 0

    for path in sorted(folder.path.glob("*.xlsx")):
        if path.name.startswith("~$"):
            continue
        kind, counts, delivered = scan_workbook(path)
        if kind == "base":
            base_counts.update(counts)
        elif kind in delivered_by_type:
            task_file_count += 1
            delivered_by_type[kind].extend(delivered)
            marker = store_group_marker_from_name(path.name)
            if marker in store_group_delivered:
                store_group_delivered[marker].extend(delivered)

    missing_chatids: set[str] = set()
    super_reach = 0
    for chat_id in delivered_by_type["super_send"]:
        if chat_id in base_counts:
            super_reach += base_counts[chat_id]
        else:
            missing_chatids.add(chat_id)

    group_reach = 0
    for chat_id in delivered_by_type["group_send"]:
        if chat_id in base_counts:
            group_reach += base_counts[chat_id]
        else:
            missing_chatids.add(chat_id)

    group_reach_row = GroupReachRow(
        target_date=folder.target_date,
        super_reach=super_reach,
        group_reach=group_reach,
        super_rows=len(delivered_by_type["super_send"]),
        group_rows=len(delivered_by_type["group_send"]),
        missing_chatid_count=len(missing_chatids),
        task_file_count=task_file_count,
        updated_at=datetime.now().replace(microsecond=0),
        folder_name=folder.path.name,
    )
    store_rows: List[StoreGroupRow] = []
    food_reach = 0
    for marker, label in STORE_GROUP_MARKERS:
        reach = sum(base_counts.get(chat_id, 0) for chat_id in store_group_delivered[marker])
        food_reach += reach
        store_rows.append(
            StoreGroupRow(
                target_date=folder.target_date,
                report_type=STORE_GROUP_TYPE,
                store_group=label,
                coupon_type=ELASTIC_COUPON_TYPE,
                reach=reach,
            )
        )
    store_rows.append(
        StoreGroupRow(
            target_date=folder.target_date,
            report_type=STORE_GROUP_TYPE,
            store_group="-",
            coupon_type=FOOD_COUPON_TYPE,
            reach=food_reach,
        )
    )

    return DailyCalculation(group_reach=group_reach_row, store_group_rows=store_rows)


def parse_docx_heading(text: str) -> Optional[date]:
    match = DATE_HEADING_RE.fullmatch(text.strip())
    if not match:
        return None
    try:
        return date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
    except ValueError:
        return None


def reach_count_from_line(text: str) -> int:
    total_match = TOTAL_REACH_RE.search(text)
    if total_match:
        return to_int(total_match.group(1))
    return sum(to_int(match.group(1)) for match in TASK_REACH_RE.finditer(text))


def read_reach_docx_totals(path: Path) -> Dict[date, Tuple[int, int]]:
    if not path.exists():
        raise ReachSummaryError(f"Reach summary docx does not exist: {path}")

    document = Document(path)
    result: Dict[date, Tuple[int, int]] = {}
    current_date: Optional[date] = None
    group_count = 0
    moment_count = 0

    def flush_current() -> None:
        if current_date is not None:
            result[current_date] = (group_count, moment_count)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        parsed = parse_docx_heading(text)
        if parsed:
            flush_current()
            current_date = parsed
            group_count = 0
            moment_count = 0
            continue
        if current_date is None:
            continue
        if GROUP_SEND_LABEL in text:
            group_count = reach_count_from_line(text)
        elif MOMENT_SEND_LABEL in text:
            moment_count = reach_count_from_line(text)

    flush_current()
    return result


def parse_date_cell(value: Any) -> Optional[date]:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    if isinstance(value, str) and value.strip():
        text = value.strip()
        for date_format in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                return datetime.strptime(text[:10], date_format).date()
            except ValueError:
                pass
        return None
    return None


def read_existing_reach_summary_rows(path: Path) -> Dict[date, ReachSummaryRow]:
    if not path.exists():
        return {}
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return {}
    if SHEET_REACH_SUMMARY not in workbook.sheetnames:
        return {}
    worksheet = workbook[SHEET_REACH_SUMMARY]
    rows: Dict[date, ReachSummaryRow] = {}
    for row in worksheet.iter_rows(min_row=2, values_only=True):
        target_date = parse_date_cell(row[0] if row else None)
        if not target_date:
            continue
        rows[target_date] = ReachSummaryRow(
            target_date=target_date,
            welfare_friends=int(row[2] or 0) if len(row) > 2 else 0,
            community=int(row[5] or 0) if len(row) > 5 else 0,
            moments=int(row[8] or 0) if len(row) > 8 else 0,
        )
    return rows


def reach_summary_dates(path: Path) -> List[date]:
    return sorted(read_existing_reach_summary_rows(path))


def store_group_dates(path: Path) -> List[date]:
    return sorted(read_existing_store_group_rows(path))


def read_existing_store_group_rows(path: Path) -> Dict[date, List[StoreGroupRow]]:
    if not path.exists():
        return {}
    try:
        workbook = load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return {}
    if SHEET_STORE_GROUP not in workbook.sheetnames:
        return {}
    worksheet = workbook[SHEET_STORE_GROUP]
    rows: Dict[date, List[StoreGroupRow]] = {}
    for row in worksheet.iter_rows(min_row=2, values_only=True):
        target_date = parse_date_cell(row[0] if row else None)
        if not target_date:
            continue
        rows.setdefault(target_date, []).append(
            StoreGroupRow(
                target_date=target_date,
                report_type=str(row[1] or STORE_GROUP_TYPE),
                store_group=str(row[2] or ""),
                coupon_type=str(row[3] or ""),
                reach=int(row[4] or 0),
            )
        )
    return rows


def ensure_output_writable(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        return
    try:
        with path.open("r+b"):
            pass
    except PermissionError as exc:
        raise ReachSummaryError(
            f"Output workbook is open or locked: {path}. Close it in Excel and rerun, "
            "or pass --output-xlsx to save to another file."
        ) from exc


def select_folders(
    folders: List[DailyFolder],
    specific_dates: Iterable[date],
    recent_days: int,
    recompute_all: bool,
) -> List[DailyFolder]:
    if recompute_all:
        return folders
    specific = set(specific_dates)
    if specific:
        return [folder for folder in folders if folder.target_date in specific]
    return folders[-max(1, recent_days) :]


def style_header(worksheet: Any) -> None:
    header_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(bold=True, color="FFFFFF")
    for cell in worksheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")


def write_workbook(
    path: Path,
    reach_rows: Sequence[ReachSummaryRow],
    store_group_rows: Sequence[StoreGroupRow],
    source_root: Path,
    recent_days: int,
) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = SHEET_REACH_SUMMARY
    worksheet.append(REACH_SUMMARY_HEADERS)
    for row in sorted(reach_rows, key=lambda item: item.target_date):
        worksheet.append(row.as_values())

    style_header(worksheet)

    for row in worksheet.iter_rows(min_row=2):
        row[0].number_format = "yyyy/mm/dd"
        row[2].number_format = "#,##0"
        row[5].number_format = "#,##0"
        row[8].number_format = "#,##0"

    worksheet.freeze_panes = "A2"
    widths = [13, 10, 16, 4, 4, 14, 4, 4, 14]
    for index, width in enumerate(widths, start=1):
        worksheet.column_dimensions[get_column_letter(index)].width = width

    store_sheet = workbook.create_sheet(SHEET_STORE_GROUP)
    store_sheet.append(STORE_GROUP_HEADERS)
    store_group_order = {label: index for index, (_marker, label) in enumerate(STORE_GROUP_MARKERS)}
    for row in sorted(
        store_group_rows,
        key=lambda item: (
            item.target_date,
            store_group_order.get(item.store_group, len(STORE_GROUP_MARKERS)),
            item.coupon_type,
        ),
    ):
        store_sheet.append(row.as_values())

    style_header(store_sheet)
    for row in store_sheet.iter_rows(min_row=2):
        row[0].number_format = "yyyy/m/d"
        row[4].number_format = "#,##0"
    store_sheet.freeze_panes = "A2"
    widths = [13, 14, 14, 12, 14]
    for index, width in enumerate(widths, start=1):
        store_sheet.column_dimensions[get_column_letter(index)].width = width

    if store_sheet.max_row >= 2:
        table_ref = f"A1:E{store_sheet.max_row}"
        table = Table(displayName="DailyStoreGroupReach", ref=table_ref)
        style = TableStyleInfo(
            name="TableStyleMedium2",
            showFirstColumn=False,
            showLastColumn=False,
            showRowStripes=True,
            showColumnStripes=False,
        )
        table.tableStyleInfo = style
        store_sheet.add_table(table)

    readme = workbook.create_sheet(SHEET_README)
    notes = [
        ["SourceRoot", str(source_root)],
        ["DefaultRecentDays", recent_days],
        ["Rule", "Delivered rows only; match task chatid to customer group ID; sum group customer totals."],
        ["ReachSummarySheet", "A/B/C/F/I columns are Date, Weekday, welfare officer friends, community, and moments."],
        ["StoreGroupSheet", "Super-send and customer-group send files with AFD/BFD/CFD/SFD in the filename are included."],
        ["StoreGroupFoodCoupon", "Food coupon reach is the sum of A/B/C/S store-group rows for the date."],
        ["Default behavior", "Recompute only latest N date folders and keep older rows from this workbook."],
        ["Full rebuild", "Run with --all to recompute every date folder."],
    ]
    for note in notes:
        readme.append(note)
    readme.column_dimensions["A"].width = 22
    readme.column_dimensions["B"].width = 120

    path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(path)


def daily_folder_map(source_root: Path, target_dates: Iterable[date], today: date) -> Dict[date, DailyFolder]:
    wanted = set(target_dates)
    available = {folder.target_date: folder for folder in find_daily_folders(source_root, FOLDER_PREFIX, today)}
    missing = sorted(wanted - set(available))
    if missing:
        formatted = ", ".join(item.isoformat() for item in missing)
        raise ReachSummaryError(f"No matching daily folder found for: {formatted}")
    return {target_date: available[target_date] for target_date in sorted(wanted)}


def selected_daily_folders(
    source_root: Path,
    today: date,
    requested_dates: Sequence[date],
    recent_days: int,
    recompute_all: bool,
) -> List[DailyFolder]:
    all_folders = find_daily_folders(source_root, FOLDER_PREFIX, today)
    selected = select_folders(all_folders, requested_dates, recent_days, recompute_all)
    if not selected:
        raise ReachSummaryError("No matching daily folders found.")
    return selected


def save_with_rows(
    output_xlsx: Path,
    source_root: Path,
    recent_days: int,
    reach_rows_by_date: Dict[date, ReachSummaryRow],
    store_rows_by_date: Dict[date, List[StoreGroupRow]],
) -> None:
    write_workbook(
        output_xlsx,
        [reach_rows_by_date[key] for key in sorted(reach_rows_by_date)],
        [row for key in sorted(store_rows_by_date) for row in store_rows_by_date[key]],
        source_root,
        recent_days,
    )


def update_reach_summary_sheet(
    source_root: Path,
    output_xlsx: Path,
    reach_docx: Path,
    target_dates: Sequence[date],
    recent_days: int = 5,
    today: Optional[date] = None,
) -> None:
    today = today or date.today()
    ensure_output_writable(output_xlsx)
    folder_by_date = daily_folder_map(source_root, target_dates, today)
    docx_totals = read_reach_docx_totals(reach_docx)
    reach_rows_by_date = read_existing_reach_summary_rows(output_xlsx)
    store_rows_by_date = read_existing_store_group_rows(output_xlsx)

    for target_date in sorted(target_dates):
        if target_date not in docx_totals:
            raise ReachSummaryError(f"No reach summary docx block found for {target_date:%Y-%m-%d}: {reach_docx}")
        welfare_friends, moments = docx_totals[target_date]
        calculation = calculate_daily_folder(folder_by_date[target_date])
        reach_rows_by_date[target_date] = ReachSummaryRow(
            target_date=target_date,
            welfare_friends=welfare_friends,
            community=calculation.group_reach.total_reach,
            moments=moments,
        )

    save_with_rows(output_xlsx, source_root, recent_days, reach_rows_by_date, store_rows_by_date)


def update_store_group_sheet(
    source_root: Path,
    output_xlsx: Path,
    target_dates: Sequence[date],
    recent_days: int = 5,
    today: Optional[date] = None,
) -> None:
    today = today or date.today()
    ensure_output_writable(output_xlsx)
    folder_by_date = daily_folder_map(source_root, target_dates, today)
    reach_rows_by_date = read_existing_reach_summary_rows(output_xlsx)
    store_rows_by_date = read_existing_store_group_rows(output_xlsx)

    for target_date in sorted(target_dates):
        calculation = calculate_daily_folder(folder_by_date[target_date])
        store_rows_by_date[target_date] = calculation.store_group_rows

    save_with_rows(output_xlsx, source_root, recent_days, reach_rows_by_date, store_rows_by_date)


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv)
    if args.summary_only and args.store_group_only:
        raise ReachSummaryError("--summary-only and --store-group-only cannot be used together.")
    today = parse_today(args.today)
    source_root = Path(args.source_root).expanduser() if args.source_root else default_source_root()
    source_root = source_root.resolve()
    output_xlsx = Path(args.output_xlsx).expanduser() if args.output_xlsx else source_root / DEFAULT_OUTPUT_NAME
    output_xlsx = output_xlsx.resolve()
    reach_docx = Path(args.reach_docx).expanduser() if args.reach_docx else source_root / DEFAULT_REACH_DOCX_NAME
    reach_docx = reach_docx.resolve()

    requested_dates = [datetime.strptime(value, "%Y-%m-%d").date() for value in args.date]
    all_folders = find_daily_folders(source_root, args.folder_prefix, today)
    selected = select_folders(all_folders, requested_dates, args.recent_days, args.all)
    if not selected:
        raise ReachSummaryError("No matching daily folders found.")
    selected_dates = [folder.target_date for folder in selected]

    if not args.dry_run:
        if not args.store_group_only:
            update_reach_summary_sheet(source_root, output_xlsx, reach_docx, selected_dates, args.recent_days, today)
        if not args.summary_only:
            update_store_group_sheet(source_root, output_xlsx, selected_dates, args.recent_days, today)
        print(f"Saved: {output_xlsx}")
        print(f"Rows recomputed: {len(selected_dates)}")
        return 0

    docx_totals = {} if args.store_group_only else read_reach_docx_totals(reach_docx)

    started = datetime.now()
    for index, folder in enumerate(selected, start=1):
        print(f"[{index}/{len(selected)}] {folder.target_date:%Y-%m-%d} {folder.path.name}", flush=True)
        calculation = calculate_daily_folder(folder)
        if not args.store_group_only:
            welfare_friends, moments = docx_totals.get(folder.target_date, (0, 0))
            print(
                f"{folder.target_date:%Y-%m-%d}\t"
                f"福利官好友={welfare_friends}\t"
                f"社群={calculation.group_reach.total_reach}\t"
                f"朋友圈={moments}\t"
                f"missing={calculation.group_reach.missing_chatid_count}"
            )
        if not args.summary_only:
            print("")
            print(SHEET_STORE_GROUP)
            for row in calculation.store_group_rows:
                print(
                    f"{row.target_date:%Y-%m-%d}\t{row.report_type}\t"
                    f"{row.store_group}\t{row.coupon_type}\t{row.reach}"
                )

    elapsed = (datetime.now() - started).total_seconds()
    print(f"Rows recomputed: {len(selected)}; elapsed: {elapsed:.1f}s")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except KeyboardInterrupt:
        print("Interrupted.", file=sys.stderr)
        raise SystemExit(130)
    except ReachSummaryError as exc:
        print(f"Error: {exc}", file=sys.stderr)
        raise SystemExit(1)
