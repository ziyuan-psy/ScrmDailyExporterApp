#!/usr/bin/env python3
"""
Summarize delivered customers for SCRM group-send and moment-send tasks.

Secrets are read from .env or environment variables. Do not hard-code tokens.
"""

from __future__ import annotations

import argparse
import re
import shutil
import sys
import time
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph

import export_super_group_undelivered as exporter
import runtime_paths


DEFAULT_OUTPUT_DOCX = "企微社群任务触达客户统计.docx"
DEFAULT_EXCLUDE_KEYWORDS = ["测试", "海外", "境外"]
DEFAULT_PAGE_SIZE = 100
DEFAULT_DOCX_LOCK_WAIT_SECONDS = 300
DEFAULT_DOCX_BACKUP_ENABLED = False
GROUP_SEND_LABEL = "群发客户"
MOMENT_SEND_LABEL = "群发朋友圈"


class ReachSummaryError(RuntimeError):
    pass


@dataclass
class ReachTask:
    template_id: str
    name: str
    send_time: str
    delivered_count: int
    raw: Dict[str, Any]


@dataclass
class ReachSummary:
    target_date: date
    group_tasks: List[ReachTask]
    moment_tasks: List[ReachTask]

    @property
    def lines(self) -> List[str]:
        return [
            format_date_heading(self.target_date),
            format_summary_line(self.target_date, GROUP_SEND_LABEL, self.group_tasks),
            format_summary_line(self.target_date, MOMENT_SEND_LABEL, self.moment_tasks),
        ]


def parse_args(argv: Sequence[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Create SCRM delivered-customer text summary for group-send and moment-send."
    )
    parser.add_argument("--date", help="Target date, YYYY-MM-DD. Defaults to yesterday.")
    parser.add_argument(
        "--output-docx",
        default=DEFAULT_OUTPUT_DOCX,
        help=f"Summary docx path. Defaults to {DEFAULT_OUTPUT_DOCX}.",
    )
    parser.add_argument("--dry-run", action="store_true", help="Print summary without editing docx.")
    return parser.parse_args(list(argv))


def yesterday() -> date:
    return date.today() - timedelta(days=1)


def parse_target_date(value: Optional[str]) -> date:
    if not value:
        return yesterday()
    return datetime.strptime(value, "%Y-%m-%d").date()


def split_keywords(value: str) -> List[str]:
    if not value:
        return []
    return [item.strip() for item in re.split(r"[,，;\n；]+", value) if item.strip()]


def env_int(env_file: Dict[str, str], key: str, default: int) -> int:
    value = exporter.env_value(env_file, key, str(default)).strip()
    try:
        return max(1, int(value))
    except ValueError:
        return default


def excluded_keywords(env_file: Dict[str, str]) -> List[str]:
    configured = split_keywords(exporter.env_value(env_file, "REACH_EXCLUDE_KEYWORDS", ""))
    return configured or DEFAULT_EXCLUDE_KEYWORDS


def should_exclude(name: str, exclude: Iterable[str]) -> bool:
    return any(keyword and keyword in name for keyword in exclude)


def to_int(value: Any) -> int:
    if value is None or value == "":
        return 0
    if isinstance(value, bool):
        return int(value)
    if isinstance(value, (int, float)):
        return int(value)
    text = str(value).strip().replace(",", "")
    if not text:
        return 0
    try:
        return int(float(text))
    except ValueError:
        return 0


def first_int_by_keys(value: Any, keys: Sequence[str]) -> Optional[int]:
    if isinstance(value, dict):
        for key in keys:
            if key in value:
                return to_int(value.get(key))
        for child in value.values():
            found = first_int_by_keys(child, keys)
            if found is not None:
                return found
    elif isinstance(value, list):
        for child in value:
            found = first_int_by_keys(child, keys)
            if found is not None:
                return found
    return None


def delivered_count_from_record(record: Dict[str, Any]) -> int:
    keys = (
        "successTargetQty",
        "deliveredCustomerCount",
        "sendCustomerCount",
        "sentCustomerCount",
        "successCustomerQty",
        "successCustomerCount",
        "customerSuccessQty",
        "arrivedCustomerQty",
    )
    value = first_int_by_keys(record, keys)
    return value if value is not None else 0


class ReachSummaryClient:
    def __init__(self, env_file: Dict[str, str]) -> None:
        self.client = exporter.ScrmClient(env_file)

    @property
    def base_url(self) -> str:
        return self.client.base_url

    def list_group_send_tasks(self, target_date: date, page_size: int) -> List[ReachTask]:
        return self._list_tasks(
            target_date=target_date,
            page_size=page_size,
            path="/bff/marketing/private/pc/groupmsg/task/list",
            task_type=1,
            referer=f"{self.base_url}/customer-marketing/message/group-send-task?type=group-send",
        )

    def list_moment_send_tasks(self, target_date: date, page_size: int) -> List[ReachTask]:
        return self._list_tasks(
            target_date=target_date,
            page_size=page_size,
            path="/bff/marketing/private/pc/groupmsg/task/template/list",
            task_type=3,
            referer=f"{self.base_url}/customer-marketing/message/group-send-task?type=moment-send",
        )

    def _list_tasks(
        self,
        target_date: date,
        page_size: int,
        path: str,
        task_type: int,
        referer: str,
    ) -> List[ReachTask]:
        start = f"{target_date:%Y-%m-%d} 00:00:00"
        end = f"{target_date:%Y-%m-%d} 23:59:59"
        page_index = 1
        tasks: List[ReachTask] = []

        while True:
            payload = {
                "taskType": task_type,
                "pageSize": page_size,
                "pageIndex": page_index,
                "sendType": None,
                "status": None,
                "templateName": "",
                "sendTimeRange": {"startTime": start, "endTime": end},
                "createUserRange": {"userIdList": [], "deptIdList": []},
            }
            data = self.client.post_json(path, payload, referer)
            page = data.get("data") or {}
            records = page.get("records") or []
            total = int(page.get("total") or len(records) or 0)
            total_pages = page.get("totalPages")

            for record in records:
                template_id = str(record.get("templateId") or "")
                if not template_id:
                    continue
                tasks.append(
                    ReachTask(
                        template_id=template_id,
                        name=str(record.get("templateName") or ""),
                        send_time=str(record.get("sendTime") or ""),
                        delivered_count=delivered_count_from_record(record),
                        raw=record,
                    )
                )

            if total_pages:
                if page_index >= int(total_pages):
                    break
            elif len(tasks) >= total or not records:
                break
            page_index += 1

        return tasks

    def fetch_detail(self, task: ReachTask) -> Dict[str, Any]:
        referer = f"{self.base_url}/customer-marketing/message/group-send-task/detail?id={task.template_id}"
        payload = {
            "templateId": task.template_id,
            "isGetSendRange": True,
            "isGetOverview": True,
            "isGetSentTip": True,
            "isGetMessagePreview": True,
        }
        data = self.client.post_json(
            "/bff/marketing/private/pc/groupmsg/task/detail",
            payload,
            referer,
        )
        return data.get("data") or {}

    def refresh_counts_from_detail(self, tasks: List[ReachTask]) -> None:
        for task in tasks:
            detail = self.fetch_detail(task)
            count = delivered_count_from_record(detail)
            if count:
                task.delivered_count = count


def build_summary(target_date: date, env_file: Dict[str, str]) -> ReachSummary:
    page_size = env_int(env_file, "REACH_PAGE_SIZE", DEFAULT_PAGE_SIZE)
    exclude = excluded_keywords(env_file)
    client = ReachSummaryClient(env_file)
    group_tasks = [
        task
        for task in client.list_group_send_tasks(target_date, page_size)
        if not should_exclude(task.name, exclude)
    ]
    moment_tasks = [
        task
        for task in client.list_moment_send_tasks(target_date, page_size)
        if not should_exclude(task.name, exclude)
    ]
    client.refresh_counts_from_detail(group_tasks + moment_tasks)
    return ReachSummary(target_date=target_date, group_tasks=group_tasks, moment_tasks=moment_tasks)


def format_date_heading(target_date: date) -> str:
    return f"{target_date.year}.{target_date.month}.{target_date.day}"


def format_summary_line(target_date: date, label: str, tasks: Sequence[ReachTask]) -> str:
    prefix = f"{target_date.month}月{target_date.day}日{label}：{len(tasks)}个任务"
    if not tasks:
        return prefix
    parts = [
        f"任务{index}触达{task.delivered_count:,}人"
        for index, task in enumerate(tasks, start=1)
    ]
    line = f"{prefix}，" + "，".join(parts)
    if len(tasks) > 1:
        line += f"，共触达{sum(task.delivered_count for task in tasks)}人"
    return line


def print_summary(summary: ReachSummary) -> None:
    print("\n".join(summary.lines))
    print("")
    print(f"{GROUP_SEND_LABEL}: {len(summary.group_tasks)} task(s)")
    for index, task in enumerate(summary.group_tasks, start=1):
        print(f"  {index}. {task.name} -> {task.delivered_count}")
    print(f"{MOMENT_SEND_LABEL}: {len(summary.moment_tasks)} task(s)")
    for index, task in enumerate(summary.moment_tasks, start=1):
        print(f"  {index}. {task.name} -> {task.delivered_count}")


DATE_HEADING_RE = re.compile(r"^(\d{4})\.(\d{1,2})\.(\d{1,2})$")


def parse_docx_date(text: str) -> Optional[date]:
    match = DATE_HEADING_RE.fullmatch(text.strip())
    if not match:
        return None
    try:
        return date(int(match.group(1)), int(match.group(2)), int(match.group(3)))
    except ValueError:
        return None


def paragraph_text(paragraph: Paragraph) -> str:
    return paragraph.text.strip()


def extract_docx_dates(path: Path) -> List[date]:
    if not path.exists():
        return []
    document = Document(path)
    dates: List[date] = []
    for paragraph in document.paragraphs:
        parsed = parse_docx_date(paragraph_text(paragraph))
        if parsed:
            dates.append(parsed)
    return sorted(set(dates))


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def insert_block_before(reference: Paragraph, lines: Sequence[str]) -> List[Paragraph]:
    inserted: List[Paragraph] = []
    for line in [*lines, ""]:
        inserted.append(reference.insert_paragraph_before(line))
    return inserted


def existing_date_blocks(document: DocumentType) -> List[Tuple[int, date]]:
    blocks: List[Tuple[int, date]] = []
    for index, paragraph in enumerate(document.paragraphs):
        parsed = parse_docx_date(paragraph_text(paragraph))
        if parsed:
            blocks.append((index, parsed))
    return blocks


def block_range(document: DocumentType, target_date: date) -> Optional[Tuple[int, int]]:
    blocks = existing_date_blocks(document)
    for block_index, (start, parsed) in enumerate(blocks):
        if parsed != target_date:
            continue
        end = blocks[block_index + 1][0] if block_index + 1 < len(blocks) else len(document.paragraphs)
        return start, end
    return None


def current_block_lines(document: DocumentType, target_date: date) -> List[str]:
    current_range = block_range(document, target_date)
    if not current_range:
        return []
    start, end = current_range
    return [paragraph_text(paragraph) for paragraph in document.paragraphs[start:end] if paragraph_text(paragraph)]


def update_document(document: DocumentType, summary: ReachSummary) -> bool:
    new_lines = summary.lines
    if current_block_lines(document, summary.target_date) == new_lines:
        return False

    current_range = block_range(document, summary.target_date)
    if current_range:
        start, end = current_range
        reference = document.paragraphs[start]
        insert_block_before(reference, new_lines)
        for paragraph in list(document.paragraphs[start + len(new_lines) + 1 : end + len(new_lines) + 1]):
            delete_paragraph(paragraph)
        return True

    blocks = existing_date_blocks(document)
    for index, parsed in blocks:
        if summary.target_date > parsed:
            insert_block_before(document.paragraphs[index], new_lines)
            return True

    if document.paragraphs and paragraph_text(document.paragraphs[-1]):
        document.add_paragraph("")
    for line in new_lines:
        document.add_paragraph(line)
    document.add_paragraph("")
    return True


def backup_docx(path: Path, backup_dir: Path) -> Optional[Path]:
    if not path.exists():
        return None
    backup_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = backup_dir / f"{path.stem}_{stamp}{path.suffix}"
    shutil.copy2(path, backup_path)
    return backup_path


def save_docx_with_lock_wait(document: DocumentType, output_docx: Path, wait_seconds: int) -> None:
    deadline = time.monotonic() + max(0, wait_seconds)
    next_notice = 0.0
    while True:
        try:
            document.save(output_docx)
            return
        except PermissionError as exc:
            if time.monotonic() >= deadline:
                raise ReachSummaryError(
                    f"Output docx is open or locked: {output_docx}. Close it and rerun."
                ) from exc
            if time.monotonic() >= next_notice:
                remaining = int(deadline - time.monotonic())
                print(
                    f"Output docx is open or locked: {output_docx}. "
                    f"Close it to continue; retrying for {remaining}s.",
                    flush=True,
                )
                next_notice = time.monotonic() + 30
            time.sleep(5)


def write_summary_docx(
    summary: ReachSummary,
    output_docx: Path,
    backup_dir: Path,
    lock_wait_seconds: int,
    backup_enabled: bool,
) -> bool:
    document = Document(output_docx) if output_docx.exists() else Document()
    changed = update_document(document, summary)
    if not changed:
        print(f"No docx change needed: {output_docx}")
        return False
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    backup_path = backup_docx(output_docx, backup_dir) if backup_enabled else None
    save_docx_with_lock_wait(document, output_docx, lock_wait_seconds)
    print(f"Updated: {output_docx}")
    if backup_path:
        print(f"Backup: {backup_path}")
    return True


def main(argv: Optional[Sequence[str]] = None) -> int:
    args = parse_args(argv or sys.argv[1:])
    env_file = exporter.load_dotenv(exporter.default_env_path())
    target_date = parse_target_date(args.date)
    summary = build_summary(target_date, env_file)

    if args.dry_run:
        print_summary(summary)
        return 0

    output_docx = Path(args.output_docx)
    if not output_docx.is_absolute():
        output_docx = runtime_paths.default_data_dir() / output_docx
    backup_dir = runtime_paths.state_dir(runtime_paths.default_config_dir()) / "backups"
    lock_wait_seconds = exporter.env_int(
        env_file,
        "REACH_DOCX_LOCK_WAIT_SECONDS",
        DEFAULT_DOCX_LOCK_WAIT_SECONDS,
    )
    backup_enabled = exporter.env_bool(
        env_file,
        "REACH_DOCX_BACKUP_ENABLED",
        DEFAULT_DOCX_BACKUP_ENABLED,
    )
    write_summary_docx(summary, output_docx, backup_dir, lock_wait_seconds, backup_enabled)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
